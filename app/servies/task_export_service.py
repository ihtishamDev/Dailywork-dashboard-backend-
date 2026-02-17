import io
import csv
import pandas as pd
from typing import Tuple
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

from app.db.database import get_connection

HEADERS = ["id", "priority", "work_needed", "phone_number", "email", "notes", "status"]

def fetch_tasks():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)

    # ✅ Ascending order for export
    cursor.execute("""
        SELECT id, priority, work_needed, phone_number, email, notes, status
        FROM tasks
        ORDER BY id ASC
    """)

    tasks = cursor.fetchall()
    cursor.close()
    conn.close()
    return tasks


def export_csv(tasks) -> io.BytesIO:
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=HEADERS)
    writer.writeheader()
    for t in tasks:
        writer.writerow({
            "id": t.get("id", ""),
            "priority": t.get("priority", ""),
            "work_needed": t.get("work_needed", ""),
            "phone_number": t.get("phone_number", ""),
            "email": t.get("email", ""),
            "notes": t.get("notes") or "",
            "status": t.get("status", "")
        })

    byte_stream = io.BytesIO(buffer.getvalue().encode("utf-8"))
    byte_stream.seek(0)
    return byte_stream


def export_excel(tasks) -> io.BytesIO:
    df = pd.DataFrame(tasks)
    # ensure column order
    for col in HEADERS:
        if col not in df.columns:
            df[col] = ""
    df = df[HEADERS]

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Tasks")
        sheet = writer.sheets["Tasks"]
        # basic formatting (auto width)
        for col_idx, col_name in enumerate(df.columns, start=1):
            max_len = max([len(str(col_name))] + [len(str(v)) for v in df[col_name].astype(str).values[:200]])
            sheet.column_dimensions[chr(64 + col_idx)].width = min(max_len + 3, 40)

    output.seek(0)
    return output


def export_docx(tasks) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Tasks Export", level=1)
    doc.add_paragraph(f"Total Records: {len(tasks)}")

    table = doc.add_table(rows=1, cols=len(HEADERS))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(HEADERS):
        hdr_cells[i].text = h

    for t in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = str(t.get("id", ""))
        row_cells[1].text = str(t.get("priority", ""))
        row_cells[2].text = str(t.get("work_needed", ""))
        row_cells[3].text = str(t.get("phone_number", ""))
        row_cells[4].text = str(t.get("email", ""))
        row_cells[5].text = str(t.get("notes") or "")
        row_cells[6].text = str(t.get("status", ""))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_pdf(tasks) -> io.BytesIO:
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    styles = getSampleStyleSheet()

    elements = []
    elements.append(Paragraph("Tasks Export", styles["Title"]))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"Total Records: {len(tasks)}", styles["Normal"]))
    elements.append(Spacer(1, 10))

    data = [HEADERS]
    for t in tasks:
        data.append([
            str(t.get("id", "")),
            str(t.get("priority", "")),
            str(t.get("work_needed", "")),
            str(t.get("phone_number", "")),
            str(t.get("email", "")),
            str(t.get("notes") or ""),
            str(t.get("status", "")),
        ])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
    ]))

    elements.append(table)
    doc.build(elements)

    output.seek(0)
    return output


def export_tasks_file(format: str) -> Tuple[io.BytesIO, str, str]:
    fmt = (format or "csv").lower().strip()
    tasks = fetch_tasks()

    if fmt == "csv":
        return export_csv(tasks), "tasks.csv", "text/csv"

    if fmt == "xlsx":
        return export_excel(tasks), "tasks.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    if fmt == "docx":
        return export_docx(tasks), "tasks.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if fmt == "pdf":
        return export_pdf(tasks), "tasks.pdf", "application/pdf"

    # fallback
    return export_csv(tasks), "tasks.csv", "text/csv"